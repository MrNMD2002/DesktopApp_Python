#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Debug chi tiết nội dung file .docx để tìm hiểu tại sao không đọc được
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
import re

def debug_file_content(file_path):
    """Debug chi tiết nội dung file"""
    print("=" * 80)
    print("🔍 DEBUG CHI TIẾT NỘI DUNG FILE .DOCX")
    print("=" * 80)
    
    if not os.path.exists(file_path):
        print(f"❌ File không tồn tại: {file_path}")
        return
    
    try:
        doc = Document(file_path)
        print(f"📄 File: {os.path.basename(file_path)}")
        print(f"📊 Tổng paragraphs: {len(doc.paragraphs)}")
        print()
        
        # Phân tích từng paragraph
        non_empty_paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                non_empty_paragraphs.append((i+1, text))
        
        print(f"📝 Paragraphs có nội dung: {len(non_empty_paragraphs)}")
        print()
        
        # Hiển thị từng dòng với phân tích chi tiết
        print("📋 PHÂN TÍCH CHI TIẾT TỪNG DÒNG:")
        print("=" * 80)
        
        for line_num, text in non_empty_paragraphs:
            print(f"\n🔍 Dòng {line_num}: '{text}'")
            
            # Phân tích từng ký tự đầu
            if text:
                first_char = text[0]
                first_5_chars = text[:5]
                print(f"   Ký tự đầu: '{first_char}' (ASCII: {ord(first_char)})")
                print(f"   5 ký tự đầu: '{first_5_chars}'")
                
                # Kiểm tra có khoảng trắng ẩn không
                if text != text.strip():
                    print(f"   ⚠️ Có khoảng trắng thừa: '{repr(text)}'")
                
                # Kiểm tra encoding
                try:
                    text.encode('utf-8')
                    print(f"   ✅ UTF-8 encoding OK")
                except UnicodeEncodeError as e:
                    print(f"   ❌ Lỗi encoding: {e}")
            
            # Test các pattern câu hỏi
            print("   🧪 Test pattern câu hỏi:")
            
            patterns = [
                (r'^Câu\s+\d+[:.]', "Câu X:"),
                (r'^\d+[:.]', "X."),
                (r'^Q\d+[:.]', "QX:"),
                (r'^QN\s*=\s*\d+[:.]', "QN=X:"),
                (r'^Question\s+\d+[:.]', "Question X:"),
                (r'^QN\s*=\s*\d+', "QN=X"),
                (r'^Câu\s+\d+', "Câu X"),
            ]
            
            for pattern, desc in patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    print(f"     ✅ Khớp: {desc}")
                else:
                    print(f"     ❌ Không khớp: {desc}")
            
            # Test pattern đáp án
            print("   🧪 Test pattern đáp án:")
            option_patterns = [
                (r'^[A-Da-d][:.)]', "A. B. C. D."),
                (r'^[A-Da-d]\s+', "A B C D"),
            ]
            
            for pattern, desc in option_patterns:
                if re.match(pattern, text):
                    print(f"     ✅ Khớp: {desc}")
                else:
                    print(f"     ❌ Không khớp: {desc}")
            
            # Test pattern đáp án đúng
            print("   🧪 Test pattern đáp án đúng:")
            answer_patterns = [
                (r'^ANSWER[:.]\s*[A-Da-d]', "ANSWER: X"),
                (r'^Đáp án[:.]\s*[A-Da-d]', "Đáp án: X"),
                (r'^Answer[:.]\s*[A-Da-d]', "Answer: X"),
                (r'^Correct[:.]\s*[A-Da-d]', "Correct: X"),
                (r'^Đúng[:.]\s*[A-Da-d]', "Đúng: X"),
            ]
            
            for pattern, desc in answer_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    print(f"     ✅ Khớp: {desc}")
                else:
                    print(f"     ❌ Không khớp: {desc}")
        
        # Tổng kết
        print("\n" + "=" * 80)
        print("📊 TỔNG KẾT PHÂN TÍCH")
        print("=" * 80)
        
        # Đếm các loại
        questions = []
        options = []
        answers = []
        
        for line_num, text in non_empty_paragraphs:
            # Kiểm tra câu hỏi
            if any(re.match(pattern, text, re.IGNORECASE) for pattern, _ in patterns):
                questions.append((line_num, text))
            
            # Kiểm tra đáp án
            if any(re.match(pattern, text) for pattern, _ in option_patterns):
                options.append((line_num, text))
            
            # Kiểm tra đáp án đúng
            if any(re.match(pattern, text, re.IGNORECASE) for pattern, _ in answer_patterns):
                answers.append((line_num, text))
        
        print(f"❓ Câu hỏi tìm thấy: {len(questions)}")
        for line_num, text in questions:
            print(f"   Dòng {line_num}: {text[:50]}{'...' if len(text) > 50 else ''}")
        
        print(f"\n🔤 Đáp án tìm thấy: {len(options)}")
        for line_num, text in options:
            print(f"   Dòng {line_num}: {text[:50]}{'...' if len(text) > 50 else ''}")
        
        print(f"\n✅ Đáp án đúng tìm thấy: {len(answers)}")
        for line_num, text in answers:
            print(f"   Dòng {line_num}: {text[:50]}{'...' if len(text) > 50 else ''}")
        
        # Gợi ý sửa lỗi
        if len(questions) == 0:
            print("\n💡 GỢI Ý KHẮC PHỤC:")
            print("1. Kiểm tra định dạng câu hỏi:")
            print("   - Phải bắt đầu bằng: QN=1:, Câu 1:, 1., Q1:, Question 1:")
            print("   - Không có khoảng trắng thừa ở đầu")
            print("   - Không có ký tự đặc biệt ẩn")
            print()
            print("2. Ví dụ định dạng đúng:")
            print("   QN=1: Nội dung câu hỏi")
            print("   Câu 1: Nội dung câu hỏi")
            print("   1. Nội dung câu hỏi")
            print("   Q1: Nội dung câu hỏi")
            print()
            print("3. Kiểm tra encoding:")
            print("   - Đảm bảo file được lưu với encoding UTF-8")
            print("   - Không có ký tự đặc biệt")
        
        return len(questions) > 0
        
    except Exception as e:
        print(f"❌ Lỗi debug file: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return False

def main():
    """Hàm main"""
    print("🚀 DEBUG CHI TIẾT FILE .DOCX")
    print("=" * 80)
    
    # Hỏi đường dẫn file
    file_path = input("Nhập đường dẫn file .docx: ").strip()
    
    if not file_path:
        print("❌ Vui lòng nhập đường dẫn file!")
        return
    
    # Debug file
    has_questions = debug_file_content(file_path)
    
    print("\n" + "=" * 80)
    if has_questions:
        print("✅ File có câu hỏi hợp lệ!")
    else:
        print("❌ File không có câu hỏi hợp lệ!")
    print("=" * 80)

if __name__ == "__main__":
    main() 