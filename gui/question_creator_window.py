import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from database.database_manager import DatabaseManager
from utils.docx_reader import DocxReader
import os

class QuestionCreatorWindow:
    def __init__(self, parent, auth_manager):
        self.parent = parent
        self.auth_manager = auth_manager
        self.db = DatabaseManager()
        self.docx_reader = DocxReader()
        self.setup_ui()
        self.load_subjects()
    
    def setup_ui(self):
        """Thiết lập giao diện người làm đề"""
        self.window = tk.Toplevel(self.parent)
        self.window.title("Người làm đề - Hệ thống Quản lý Đề thi")
        self.window.geometry("600x500")
        
        # Frame chính
        main_frame = ttk.Frame(self.window, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Header
        header_frame = ttk.Frame(main_frame)
        header_frame.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        
        user_info = self.auth_manager.get_current_user()
        ttk.Label(header_frame, text=f"Chào mừng: {user_info['full_name']}", 
                 font=("Arial", 12, "bold")).pack(side=tk.LEFT)
        
        ttk.Button(header_frame, text="Đăng xuất", 
                  command=self.logout).pack(side=tk.RIGHT)
        
        # Frame chọn môn học
        subject_frame = ttk.LabelFrame(main_frame, text="Chọn môn học", padding="10")
        subject_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        
        ttk.Label(subject_frame, text="Môn học:").grid(row=0, column=0, sticky="w")
        
        self.subject_var = tk.StringVar()
        self.subject_combo = ttk.Combobox(subject_frame, textvariable=self.subject_var, 
                                         state="readonly", width=30)
        self.subject_combo.grid(row=0, column=1, padx=(10, 0), sticky="w")
        
        # Frame upload file
        upload_frame = ttk.LabelFrame(main_frame, text="Upload file .docx", padding="10")
        upload_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        
        self.file_path_var = tk.StringVar()
        ttk.Entry(upload_frame, textvariable=self.file_path_var, width=50).grid(row=0, column=0, sticky="ew")
        
        ttk.Button(upload_frame, text="Chọn file", 
                  command=self.select_file).grid(row=0, column=1, padx=(10, 0))
        
        ttk.Button(upload_frame, text="Đọc file", 
                  command=self.read_file).grid(row=1, column=0, columnspan=2, pady=10)
        
        # Frame nút bấm
        button_frame = ttk.Frame(upload_frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=(0, 10))
        
        ttk.Button(button_frame, text="Test File", 
                  command=self.test_file).grid(row=0, column=0, padx=(0, 5))
        
        ttk.Button(button_frame, text="Tạo File Mẫu", 
                  command=self.create_sample_file).grid(row=0, column=1, padx=5)
        
        # Frame hướng dẫn
        guide_frame = ttk.LabelFrame(main_frame, text="Hướng dẫn định dạng", padding="10")
        guide_frame.grid(row=3, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        
        guide_text = self.docx_reader.get_template_instructions()
        guide_label = ttk.Label(guide_frame, text=guide_text, justify=tk.LEFT)
        guide_label.grid(row=0, column=0, sticky="w")
        
        # Frame thống kê
        stats_frame = ttk.LabelFrame(main_frame, text="Thống kê câu hỏi", padding="10")
        stats_frame.grid(row=4, column=0, columnspan=2, sticky="ew")
        
        # Treeview cho thống kê
        columns = ("Môn học", "Tổng câu hỏi", "Dễ", "Trung bình", "Khó")
        self.stats_tree = ttk.Treeview(stats_frame, columns=columns, show="headings", height=5)
        
        for col in columns:
            self.stats_tree.heading(col, text=col)
            self.stats_tree.column(col, width=100)
        
        self.stats_tree.grid(row=0, column=0, sticky="ew")
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(stats_frame, orient="vertical", command=self.stats_tree.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.stats_tree.configure(yscrollcommand=scrollbar.set)
        
        # Nút refresh thống kê
        ttk.Button(stats_frame, text="Làm mới thống kê", 
                  command=self.load_statistics).grid(row=1, column=0, pady=10)
        
        # Cấu hình grid
        self.window.columnconfigure(0, weight=1)
        self.window.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        upload_frame.columnconfigure(0, weight=1)
        stats_frame.columnconfigure(0, weight=1)
    
    def load_subjects(self):
        """Tải danh sách môn học"""
        try:
            query = "SELECT id, name FROM subjects ORDER BY name"
            subjects = self.db.execute_query(query)
            
            subject_dict = {}
            subject_names = []
            
            for subject in subjects:
                subject_dict[subject['name']] = subject['id']
                subject_names.append(subject['name'])
            
            self.subject_combo['values'] = subject_names
            self.subject_dict = subject_dict
            
            if subject_names:
                self.subject_combo.set(subject_names[0])
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể tải danh sách môn học: {str(e)}")
    
    def select_file(self):
        """Chọn file .docx"""
        file_path = filedialog.askopenfilename(
            title="Chọn file .docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        
        if file_path:
            self.file_path_var.set(file_path)
    
    def read_file(self):
        """Đọc file .docx và import câu hỏi"""
        file_path = self.file_path_var.get().strip()
        subject_name = self.subject_var.get()
        
        if not file_path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn file!")
            return
        
        if not subject_name:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn môn học!")
            return
        
        # Kiểm tra file tồn tại
        if not os.path.exists(file_path):
            messagebox.showerror("Lỗi", f"File không tồn tại:\n{file_path}")
            return
        
        # Kiểm tra định dạng file
        if not file_path.lower().endswith('.docx'):
            messagebox.showerror("Lỗi", f"File không phải định dạng .docx:\n{file_path}")
            return
        
        subject_id = self.subject_dict.get(subject_name)
        creator_id = self.auth_manager.get_current_user()['id']
        
        try:
            # Hiển thị dialog xác nhận
            result = messagebox.askyesno("Xác nhận", 
                                       f"Bạn có muốn import câu hỏi từ file này vào môn {subject_name}?\n\nFile: {os.path.basename(file_path)}")
            
            if result:
                # Hiển thị thông báo đang xử lý
                processing_window = tk.Toplevel(self.window)
                processing_window.title("Đang xử lý...")
                processing_window.geometry("400x150")
                processing_window.transient(self.window)
                processing_window.grab_set()
                
                # Căn giữa cửa sổ
                processing_window.update_idletasks()
                x = (processing_window.winfo_screenwidth() // 2) - (200)
                y = (processing_window.winfo_screenheight() // 2) - (75)
                processing_window.geometry(f'400x150+{x}+{y}')
                
                # Nội dung
                ttk.Label(processing_window, text="Đang đọc file...", 
                          font=("Arial", 12)).pack(pady=20)
                
                progress = ttk.Progressbar(processing_window, mode='indeterminate')
                progress.pack(pady=10, padx=20, fill='x')
                progress.start()
                
                # Cập nhật giao diện
                processing_window.update()
                
                try:
                    # Đọc file
                    success, message = self.docx_reader.read_docx_file(file_path, subject_id, creator_id)
                    
                    # Đóng cửa sổ xử lý
                    processing_window.destroy()
                    
                    if success:
                        messagebox.showinfo("Thành công", message)
                        self.file_path_var.set("")  # Xóa đường dẫn file
                        self.load_statistics()  # Cập nhật thống kê
                    else:
                        # Hiển thị lỗi chi tiết
                        error_details = f"Lỗi đọc file:\n\n{message}\n\n"
                        error_details += "🔧 Hướng dẫn khắc phục:\n"
                        error_details += "1. Kiểm tra định dạng file theo hướng dẫn\n"
                        error_details += "2. Đảm bảo file không bị hỏng\n"
                        error_details += "3. Kiểm tra quyền đọc file\n"
                        error_details += "4. Chạy test: python test_file_reading.py"
                        
                        messagebox.showerror("Lỗi", error_details)
                
                except Exception as e:
                    # Đóng cửa sổ xử lý
                    processing_window.destroy()
                    
                    # Hiển thị lỗi chi tiết
                    error_details = f"Lỗi không mong muốn:\n\n{str(e)}\n\n"
                    error_details += "🔧 Hướng dẫn khắc phục:\n"
                    error_details += "1. Kiểm tra file có hợp lệ không\n"
                    error_details += "2. Kiểm tra kết nối database\n"
                    error_details += "3. Chạy test: python test_file_reading.py\n"
                    error_details += "4. Kiểm tra log file: app.log"
                    
                    messagebox.showerror("Lỗi", error_details)
        
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file: {str(e)}")
    
    def test_file(self):
        """Test file trước khi import - chi tiết từng dòng"""
        file_path = self.file_path_var.get().strip()
        
        if not file_path:
            messagebox.showwarning("Cảnh báo", "Vui lòng chọn file trước khi test!")
            return
        
        # Kiểm tra file tồn tại
        if not os.path.exists(file_path):
            messagebox.showerror("Lỗi", f"File không tồn tại:\n{file_path}")
            return
        
        # Kiểm tra định dạng file
        if not file_path.lower().endswith('.docx'):
            messagebox.showerror("Lỗi", f"File không phải định dạng .docx:\n{file_path}")
            return
        
        try:
            # Sử dụng hàm test chi tiết mới
            success, result = self.docx_reader.test_file_detailed(file_path)
            
            # Hiển thị dialog kết quả
            dialog = tk.Toplevel(self.window)
            dialog.title("Kết quả Test File")
            dialog.geometry("600x500")
            dialog.transient(self.window)
            dialog.grab_set()
            
            # Căn giữa cửa sổ
            dialog.update_idletasks()
            x = (dialog.winfo_screenwidth() // 2) - (300)
            y = (dialog.winfo_screenheight() // 2) - (250)
            dialog.geometry(f'600x500+{x}+{y}')
            
            # Text widget
            text_widget = tk.Text(dialog, wrap="word", padx=10, pady=10, font=("Consolas", 9))
            text_widget.pack(fill="both", expand=True, padx=10, pady=10)
            
            # Scrollbar
            scrollbar = ttk.Scrollbar(dialog, orient="vertical", command=text_widget.yview)
            scrollbar.pack(side="right", fill="y")
            text_widget.configure(yscrollcommand=scrollbar.set)
            
            # Hiển thị kết quả
            text_widget.insert("1.0", result)
            text_widget.config(state="disabled")
            
            # Nút đóng
            ttk.Button(dialog, text="Close", command=dialog.destroy).pack(pady=10)
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể test file: {str(e)}")
    
    def create_sample_file(self):
        """Tạo file mẫu"""
        try:
            from docx import Document
            
            # Tạo file mẫu
            doc = Document()
            
            # Thêm tiêu đề
            doc.add_paragraph("SAMPLE QUESTION FILE")
            doc.add_paragraph("Format: QN=X: Question")
            doc.add_paragraph("")
            
            # Câu hỏi 1
            doc.add_paragraph("QN=1: What is the capital of Vietnam?")
            doc.add_paragraph("a. Hanoi")
            doc.add_paragraph("b. Ho Chi Minh City")
            doc.add_paragraph("c. Da Nang")
            doc.add_paragraph("d. Hue")
            doc.add_paragraph("ANSWER: A")
            doc.add_paragraph("MARK: 1.0")
            doc.add_paragraph("UNIT: Geography")
            doc.add_paragraph("")
            
            # Câu hỏi 2
            doc.add_paragraph("Câu 2: 2 + 2 = ?")
            doc.add_paragraph("A. 3")
            doc.add_paragraph("B. 4")
            doc.add_paragraph("C. 5")
            doc.add_paragraph("D. 6")
            doc.add_paragraph("Đáp án: B")
            doc.add_paragraph("Điểm: 0.5")
            doc.add_paragraph("Đơn vị: Toán học")
            
            # Lưu file
            sample_file = "sample_questions.docx"
            doc.save(sample_file)
            
            messagebox.showinfo("Thành công", 
                              f"Đã tạo file mẫu: {sample_file}\n\n"
                              "Bạn có thể sử dụng file này để test chức năng import.")
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể tạo file mẫu: {str(e)}")
    
    def load_statistics(self):
        """Tải thống kê câu hỏi"""
        try:
            query = """
                SELECT s.name as subject_name,
                       COUNT(q.id) as total_questions,
                       SUM(CASE WHEN q.difficulty_level = 'easy' THEN 1 ELSE 0 END) as easy_count,
                       SUM(CASE WHEN q.difficulty_level = 'medium' THEN 1 ELSE 0 END) as medium_count,
                       SUM(CASE WHEN q.difficulty_level = 'hard' THEN 1 ELSE 0 END) as hard_count
                FROM subjects s
                LEFT JOIN questions q ON s.id = q.subject_id
                GROUP BY s.id, s.name
                ORDER BY s.name
            """
            
            stats = self.db.execute_query(query)
            
            # Xóa dữ liệu cũ
            for item in self.stats_tree.get_children():
                self.stats_tree.delete(item)
            
            # Thêm dữ liệu mới
            for stat in stats:
                self.stats_tree.insert("", "end", values=(
                    stat['subject_name'],
                    stat['total_questions'] or 0,
                    stat['easy_count'] or 0,
                    stat['medium_count'] or 0,
                    stat['hard_count'] or 0
                ))
        
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể tải thống kê: {str(e)}")
    
    def logout(self):
        """Đăng xuất"""
        self.auth_manager.logout()
        self.window.destroy()
        messagebox.showinfo("Thông báo", "Đã đăng xuất thành công!") 