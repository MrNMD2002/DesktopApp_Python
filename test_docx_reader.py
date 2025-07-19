#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Demo chức năng đọc file Word với định dạng mới
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from utils.docx_reader import DocxReader
from docx import Document
import tempfile

def create_test_docx():
    """Tạo file Word test với định dạng như trong hình"""
    doc = Document()
    
    # Thêm tiêu đề
    doc.add_paragraph("Subject: ISC")
    doc.add_paragraph("Number of Quiz: 30")
    doc.add_paragraph("Lecturer: hungpd2")
    doc.add_paragraph("Date: dd-mm-yyyy")
    doc.add_paragraph("")  # Dòng trống
    
    # Câu hỏi 1
    doc.add_paragraph("QN=1: See the figure and choose the right type of B2B E-Commerce [file:8435.jpg]")
    doc.add_paragraph("a. Sell-side B2B")
    doc.add_paragraph("b. Electronic Exchange")
    doc.add_paragraph("c. Buy-side B2B")
    doc.add_paragraph("d. Supply Chain Improvements and Collaborative Commerce")
    doc.add_paragraph("ANSWER: B")
    doc.add_paragraph("MARK: 0.5")
    doc.add_paragraph("UNIT: Chapter1")
    doc.add_paragraph("MIX CHOICES: Yes")
    doc.add_paragraph("")  # Dòng trống
    
    # Câu hỏi 2
    doc.add_paragraph("QN=2: What is the main advantage of B2B E-Commerce?")
    doc.add_paragraph("a. Lower costs for consumers")
    doc.add_paragraph("b. Reduced transaction costs")
    doc.add_paragraph("c. Faster delivery times")
    doc.add_paragraph("d. Better customer service")
    doc.add_paragraph("ANSWER: B")
    doc.add_paragraph("MARK: 1.0")
    doc.add_paragraph("UNIT: Chapter1")
    doc.add_paragraph("MIX CHOICES: No")
    doc.add_paragraph("")  # Dòng trống
    
    # Câu hỏi 3 (định dạng khác)
    doc.add_paragraph("Câu 3: Thủ đô của Việt Nam là?")
    doc.add_paragraph("A. Hà Nội")
    doc.add_paragraph("B. TP. Hồ Chí Minh")
    doc.add_paragraph("C. Đà Nẵng")
    doc.add_paragraph("D. Huế")
    doc.add_paragraph("Đáp án: A")
    doc.add_paragraph("Điểm: 1.0")
    doc.add_paragraph("Đơn vị: Địa lý")
    doc.add_paragraph("")  # Dòng trống
    
    # Câu hỏi 4 (định dạng khác)
    doc.add_paragraph("4. 2 + 2 = ?")
    doc.add_paragraph("a) 3")
    doc.add_paragraph("b) 4")
    doc.add_paragraph("c) 5")
    doc.add_paragraph("d) 6")
    doc.add_paragraph("Answer: B")
    doc.add_paragraph("Mark: 0.5")
    doc.add_paragraph("Unit: Toán học")
    
    return doc

def demo_docx_reading():
    """Demo đọc file Word"""
    print("=" * 60)
    print("DEMO ĐỌC FILE WORD VỚI ĐỊNH DẠNG MỚI")
    print("=" * 60)
    
    try:
        # Tạo file Word test
        print("📝 Tạo file Word test...")
        doc = create_test_docx()
        
        # Lưu file tạm thời
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            test_file_path = tmp_file.name
        
        print(f"✅ File test đã tạo: {test_file_path}")
        
        # Đọc file
        print("\n📖 Đọc file Word...")
        docx_reader = DocxReader()
        
        # Test với subject_id và creator_id giả
        success, message = docx_reader.read_docx_file(test_file_path, subject_id=1, creator_id=1)
        
        if success:
            print(f"✅ {message}")
        else:
            print(f"❌ {message}")
        
        # Xóa file tạm
        os.unlink(test_file_path)
        print(f"🗑️ Đã xóa file test: {test_file_path}")
        
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()

def demo_template_instructions():
    """Demo hướng dẫn template"""
    print("\n" + "=" * 60)
    print("HƯỚNG DẪN ĐỊNH DẠNG TEMPLATE")
    print("=" * 60)
    
    docx_reader = DocxReader()
    instructions = docx_reader.get_template_instructions()
    print(instructions)

def demo_parsing_functions():
    """Demo các hàm parse riêng lẻ"""
    print("\n" + "=" * 60)
    print("DEMO CÁC HÀM PARSE")
    print("=" * 60)
    
    docx_reader = DocxReader()
    
    # Test parse câu hỏi
    test_questions = [
        "QN=1: Test question",
        "Câu 2: Test question",
        "3. Test question",
        "Q4: Test question"
    ]
    
    print("🔍 Test parse số câu hỏi:")
    for question in test_questions:
        number = docx_reader._extract_question_number(question)
        print(f"   '{question}' -> {number}")
    
    # Test parse đáp án
    test_options = [
        "a. Option A",
        "B. Option B",
        "c) Option C",
        "D Option D"
    ]
    
    print("\n🔍 Test parse đáp án:")
    for option in test_options:
        letter, text = docx_reader._extract_option(option)
        print(f"   '{option}' -> {letter}: {text}")
    
    # Test parse đáp án đúng
    test_answers = [
        "ANSWER: A",
        "Đáp án: B",
        "Answer: C",
        "Correct: D"
    ]
    
    print("\n🔍 Test parse đáp án đúng:")
    for answer in test_answers:
        correct = docx_reader._extract_correct_answer(answer)
        print(f"   '{answer}' -> {correct}")
    
    # Test parse thông tin bổ sung
    test_mark = "MARK: 0.5"
    test_unit = "UNIT: Chapter1"
    test_mix = "MIX CHOICES: Yes"
    
    print("\n🔍 Test parse thông tin bổ sung:")
    print(f"   '{test_mark}' -> {docx_reader._extract_mark(test_mark)}")
    print(f"   '{test_unit}' -> {docx_reader._extract_unit(test_unit)}")
    print(f"   '{test_mix}' -> {docx_reader._extract_mix_choices(test_mix)}")

def main():
    """Hàm main"""
    print("🚀 DEMO CHỨC NĂNG ĐỌC FILE WORD")
    print("=" * 60)
    
    try:
        # Demo parse functions
        demo_parsing_functions()
        
        # Demo template instructions
        demo_template_instructions()
        
        # Demo đọc file Word
        demo_docx_reading()
        
        print("\n" + "=" * 60)
        print("🎉 DEMO HOÀN THÀNH!")
        print("=" * 60)
        print("Bạn có thể tạo file Word theo định dạng trên và test trong ứng dụng chính")
        print("python main.py")
        
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 