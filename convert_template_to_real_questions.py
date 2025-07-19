#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Công cụ chuyển đổi template placeholder thành câu hỏi thực tế
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches
import re

def analyze_template_structure(file_path):
    """Phân tích cấu trúc template"""
    print("=" * 80)
    print("🔍 PHÂN TÍCH CẤU TRÚC TEMPLATE")
    print("=" * 80)
    
    try:
        doc = Document(file_path)
        print(f"📄 File: {os.path.basename(file_path)}")
        print(f"📊 Tổng paragraphs: {len(doc.paragraphs)}")
        print()
        
        # Tìm các câu hỏi placeholder
        questions = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text.startswith('QN='):
                questions.append((i+1, text))
        
        print(f"❓ Tìm thấy {len(questions)} câu hỏi placeholder:")
        for line_num, text in questions:
            print(f"   Dòng {line_num}: {text}")
        
        return len(questions)
        
    except Exception as e:
        print(f"❌ Lỗi phân tích: {e}")
        return 0

def create_real_questions_file():
    """Tạo file với câu hỏi thực tế và hình ảnh"""
    print("\n" + "=" * 80)
    print("📝 TẠO FILE VỚI CÂU HỎI THỰC TẾ VÀ HÌNH ẢNH")
    print("=" * 80)
    
    try:
        doc = Document()
        
        # Header
        doc.add_paragraph("QUIZ TEMPLATE - CÂU HỎI THỰC TẾ")
        doc.add_paragraph("Subject: ISC")
        doc.add_paragraph("Number of Quiz: 20")
        doc.add_paragraph("Lecturer: hungpd2")
        doc.add_paragraph("Date: dd-mm-yyyy")
        doc.add_paragraph("")
        
        # Câu hỏi 1 - B2B E-Commerce (có hình ảnh)
        doc.add_paragraph("QN=1: See the figure and choose the right type of B2B E-Commerce [file:8435.jpg]")
        doc.add_paragraph("a. Sell-side B2B")
        doc.add_paragraph("b. Electronic Exchange")
        doc.add_paragraph("c. Buy-side B2B")
        doc.add_paragraph("d. Supply Chain Improvements and Collaborative Commerce")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 2 - Toán học (có hình ảnh)
        doc.add_paragraph("QN=2: Solve the math problem shown in the image [file:math_problem.png]")
        doc.add_paragraph("a. 40")
        doc.add_paragraph("b. 42")
        doc.add_paragraph("c. 43")
        doc.add_paragraph("d. 44")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 3 - Địa lý (không có hình ảnh)
        doc.add_paragraph("QN=3: Thủ đô của Việt Nam là?")
        doc.add_paragraph("a. TP. Hồ Chí Minh")
        doc.add_paragraph("b. Hà Nội")
        doc.add_paragraph("c. Đà Nẵng")
        doc.add_paragraph("d. Huế")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 4 - Lịch sử (có hình ảnh)
        doc.add_paragraph("QN=4: Identify the historical event shown in this image [file:independence_day.jpg]")
        doc.add_paragraph("a. 1944")
        doc.add_paragraph("b. 1945")
        doc.add_paragraph("c. 1946")
        doc.add_paragraph("d. 1947")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 5 - Khoa học (không có hình ảnh)
        doc.add_paragraph("QN=5: Nguyên tố nào có ký hiệu H?")
        doc.add_paragraph("a. Helium")
        doc.add_paragraph("b. Hydrogen")
        doc.add_paragraph("c. Carbon")
        doc.add_paragraph("d. Nitrogen")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 6 - Văn học (có hình ảnh)
        doc.add_paragraph("QN=6: Look at the image and identify the author of this literary work [file:truyen_kieu.jpg]")
        doc.add_paragraph("a. Hồ Xuân Hương")
        doc.add_paragraph("b. Nguyễn Du")
        doc.add_paragraph("c. Nguyễn Trãi")
        doc.add_paragraph("d. Nguyễn Bỉnh Khiêm")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 7 - Công nghệ (không có hình ảnh)
        doc.add_paragraph("QN=7: HTML là viết tắt của?")
        doc.add_paragraph("a. Hyper Text Markup Language")
        doc.add_paragraph("b. High Tech Modern Language")
        doc.add_paragraph("c. Home Tool Markup Language")
        doc.add_paragraph("d. Hyperlink and Text Markup Language")
        doc.add_paragraph("ANSWER: A")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 8 - Thể thao (có hình ảnh)
        doc.add_paragraph("QN=8: Based on the image, which sport is shown? [file:football.jpg]")
        doc.add_paragraph("a. Bóng rổ")
        doc.add_paragraph("b. Bóng đá")
        doc.add_paragraph("c. Tennis")
        doc.add_paragraph("d. Bóng chuyền")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        
        # Lưu file
        output_file = "Quiz_Real_Questions_With_Images.docx"
        doc.save(output_file)
        
        print(f"✅ Đã tạo file với câu hỏi thực tế: {output_file}")
        print("\n📋 NỘI DUNG FILE:")
        print("- 8 câu hỏi thực tế từ nhiều lĩnh vực")
        print("- 5 câu hỏi có tham chiếu hình ảnh [file:...]")
        print("- 3 câu hỏi chỉ có text")
        print("- Định dạng đúng theo yêu cầu")
        print("- Đầy đủ đáp án và metadata")
        print("- Sẵn sàng để test import")
        
        return output_file
        
    except Exception as e:
        print(f"❌ Lỗi tạo file: {e}")
        return None

def replace_placeholders_in_file(input_file, output_file):
    """Thay thế placeholder trong file hiện tại"""
    print("\n" + "=" * 80)
    print("🔧 THAY THẾ PLACEHOLDER TRONG FILE HIỆN TẠI")
    print("=" * 80)
    
    try:
        # Đọc file hiện tại
        doc = Document(input_file)
        
        # Danh sách câu hỏi thực tế với hình ảnh
        real_questions = [
            {
                'question': 'QN=1: See the figure and choose the right type of B2B E-Commerce [file:8435.jpg]',
                'options': ['a. Sell-side B2B', 'b. Electronic Exchange', 'c. Buy-side B2B', 'd. Supply Chain Improvements and Collaborative Commerce'],
                'answer': 'ANSWER: B',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            },
            {
                'question': 'QN=2: Solve the math problem shown in the image [file:math_problem.png]',
                'options': ['a. 40', 'b. 42', 'c. 43', 'd. 44'],
                'answer': 'ANSWER: B',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            },
            {
                'question': 'QN=3: Thủ đô của Việt Nam là?',
                'options': ['a. TP. Hồ Chí Minh', 'b. Hà Nội', 'c. Đà Nẵng', 'd. Huế'],
                'answer': 'ANSWER: B',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            },
            {
                'question': 'QN=4: Identify the historical event shown in this image [file:independence_day.jpg]',
                'options': ['a. 1944', 'b. 1945', 'c. 1946', 'd. 1947'],
                'answer': 'ANSWER: B',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            },
            {
                'question': 'QN=5: Nguyên tố nào có ký hiệu H?',
                'options': ['a. Helium', 'b. Hydrogen', 'c. Carbon', 'd. Nitrogen'],
                'answer': 'ANSWER: B',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            },
            {
                'question': 'QN=6: Look at the image and identify the author of this literary work [file:truyen_kieu.jpg]',
                'options': ['a. Hồ Xuân Hương', 'b. Nguyễn Du', 'c. Nguyễn Trãi', 'd. Nguyễn Bỉnh Khiêm'],
                'answer': 'ANSWER: B',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            },
            {
                'question': 'QN=7: HTML là viết tắt của?',
                'options': ['a. Hyper Text Markup Language', 'b. High Tech Modern Language', 'c. Home Tool Markup Language', 'd. Hyperlink and Text Markup Language'],
                'answer': 'ANSWER: A',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            },
            {
                'question': 'QN=8: Based on the image, which sport is shown? [file:football.jpg]',
                'options': ['a. Bóng rổ', 'b. Bóng đá', 'c. Tennis', 'd. Bóng chuyền'],
                'answer': 'ANSWER: B',
                'mark': 'MARK: 0.5',
                'unit': 'UNIT: Chapter1',
                'mix': 'MIX CHOICES: Yes'
            }
        ]
        
        # Thay thế từng paragraph
        question_index = 0
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Tìm câu hỏi placeholder
            if text.startswith('QN=') and question_index < len(real_questions):
                # Thay thế câu hỏi
                paragraph.text = real_questions[question_index]['question']
                
                # Thay thế các dòng tiếp theo
                next_index = i + 1
                for option in real_questions[question_index]['options']:
                    if next_index < len(doc.paragraphs):
                        doc.paragraphs[next_index].text = option
                        next_index += 1
                
                if next_index < len(doc.paragraphs):
                    doc.paragraphs[next_index].text = real_questions[question_index]['answer']
                    next_index += 1
                
                if next_index < len(doc.paragraphs):
                    doc.paragraphs[next_index].text = real_questions[question_index]['mark']
                    next_index += 1
                
                if next_index < len(doc.paragraphs):
                    doc.paragraphs[next_index].text = real_questions[question_index]['unit']
                    next_index += 1
                
                if next_index < len(doc.paragraphs):
                    doc.paragraphs[next_index].text = real_questions[question_index]['mix']
                
                question_index += 1
        
        # Lưu file đã sửa
        doc.save(output_file)
        
        print(f"✅ Đã thay thế {question_index} câu hỏi placeholder")
        print(f"📄 File đã sửa: {output_file}")
        
        return output_file
        
    except Exception as e:
        print(f"❌ Lỗi thay thế: {e}")
        return None

def create_demo_with_images():
    """Tạo file demo với hình ảnh thực tế"""
    print("\n" + "=" * 80)
    print("🖼️ TẠO FILE DEMO VỚI HÌNH ẢNH THỰC TẾ")
    print("=" * 80)
    
    try:
        # Tạo thư mục images nếu chưa có
        images_dir = "demo_images"
        if not os.path.exists(images_dir):
            os.makedirs(images_dir)
            print(f"📁 Đã tạo thư mục: {images_dir}")
        
        # Tạo file demo
        doc = Document()
        
        # Header
        doc.add_paragraph("DEMO QUIZ - HÌNH ẢNH THỰC TẾ")
        doc.add_paragraph("Subject: Demo")
        doc.add_paragraph("Number of Quiz: 3")
        doc.add_paragraph("Lecturer: Demo")
        doc.add_paragraph("Date: dd-mm-yyyy")
        doc.add_paragraph("")
        
        # Câu hỏi 1 - Có hình ảnh
        doc.add_paragraph("QN=1: What type of B2B E-Commerce is shown in this diagram? [file:b2b_diagram.jpg]")
        doc.add_paragraph("a. Sell-side B2B")
        doc.add_paragraph("b. Electronic Exchange")
        doc.add_paragraph("c. Buy-side B2B")
        doc.add_paragraph("d. Supply Chain")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 1.0")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 2 - Không có hình ảnh
        doc.add_paragraph("QN=2: What is 25 + 17?")
        doc.add_paragraph("a. 40")
        doc.add_paragraph("b. 42")
        doc.add_paragraph("c. 43")
        doc.add_paragraph("d. 44")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 1.0")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 3 - Có hình ảnh
        doc.add_paragraph("QN=3: Identify the programming language shown in this code snippet [file:python_code.png]")
        doc.add_paragraph("a. Java")
        doc.add_paragraph("b. Python")
        doc.add_paragraph("c. C++")
        doc.add_paragraph("d. JavaScript")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 1.0")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        
        # Lưu file
        output_file = "Demo_Quiz_With_Images.docx"
        doc.save(output_file)
        
        print(f"✅ Đã tạo file demo: {output_file}")
        print(f"📁 Thư mục hình ảnh: {images_dir}")
        print("\n📋 HƯỚNG DẪN SỬ DỤNG:")
        print("1. Đặt hình ảnh vào thư mục 'demo_images'")
        print("2. Đặt tên file: b2b_diagram.jpg, python_code.png")
        print("3. Test import file .docx vào hệ thống")
        print("4. Hệ thống sẽ tự động trích xuất và lưu hình ảnh")
        
        return output_file
        
    except Exception as e:
        print(f"❌ Lỗi tạo demo: {e}")
        return None

def main():
    """Hàm main"""
    print("🚀 CÔNG CỤ CHUYỂN ĐỔI TEMPLATE THÀNH CÂU HỎI THỰC TẾ")
    print("=" * 80)
    
    print("\n" + "=" * 80)
    print("🔧 TÙY CHỌN")
    print("=" * 80)
    print("1. Tạo file mới với câu hỏi thực tế và hình ảnh")
    print("2. Thay thế placeholder trong file hiện tại")
    print("3. Tạo file demo với hình ảnh thực tế")
    print("4. Cả hai (tạo mới + thay thế)")
    
    choice = input("\nChọn tùy chọn (1/2/3/4): ").strip()
    
    if choice == "1":
        # Tạo file mới
        new_file = create_real_questions_file()
        if new_file:
            print(f"\n✅ Đã tạo file mới: {new_file}")
            print("💡 Sử dụng file này để test chức năng import")
    
    elif choice == "2":
        # Thay thế trong file hiện tại
        file_path = input("Nhập đường dẫn file template .docx: ").strip()
        if not file_path:
            print("❌ Vui lòng nhập đường dẫn file!")
            return
        
        question_count = analyze_template_structure(file_path)
        if question_count == 0:
            print("❌ Không tìm thấy câu hỏi placeholder!")
            return
        
        output_file = file_path.replace('.docx', '_real_questions.docx')
        fixed_file = replace_placeholders_in_file(file_path, output_file)
        if fixed_file:
            print(f"\n✅ Đã sửa file: {fixed_file}")
            print("💡 File đã được thay thế placeholder bằng câu hỏi thực tế")
    
    elif choice == "3":
        # Tạo demo với hình ảnh
        demo_file = create_demo_with_images()
        if demo_file:
            print(f"\n✅ Đã tạo file demo: {demo_file}")
            print("💡 Sử dụng file này để test chức năng hình ảnh")
    
    elif choice == "4":
        # Cả hai
        file_path = input("Nhập đường dẫn file template .docx: ").strip()
        if not file_path:
            print("❌ Vui lòng nhập đường dẫn file!")
            return
        
        question_count = analyze_template_structure(file_path)
        if question_count == 0:
            print("❌ Không tìm thấy câu hỏi placeholder!")
            return
        
        new_file = create_real_questions_file()
        output_file = file_path.replace('.docx', '_real_questions.docx')
        fixed_file = replace_placeholders_in_file(file_path, output_file)
        
        if new_file and fixed_file:
            print(f"\n✅ Đã tạo cả hai file:")
            print(f"   📄 File mới: {new_file}")
            print(f"   🔧 File đã sửa: {fixed_file}")
            print("💡 Bạn có thể sử dụng cả hai để test")
    
    else:
        print("❌ Lựa chọn không hợp lệ!")
    
    print("\n" + "=" * 80)
    print("🎉 HOÀN THÀNH")
    print("=" * 80)

if __name__ == "__main__":
    main() 