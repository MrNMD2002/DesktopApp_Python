#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Công cụ kiểm tra và sửa file template .docx
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Inches
import re

def analyze_template_file(file_path):
    """Phân tích file template hiện tại"""
    print("=" * 80)
    print("🔍 PHÂN TÍCH FILE TEMPLATE")
    print("=" * 80)
    
    if not os.path.exists(file_path):
        print(f"❌ File không tồn tại: {file_path}")
        return False
    
    try:
        doc = Document(file_path)
        print(f"📄 File: {os.path.basename(file_path)}")
        print(f"📊 Tổng paragraphs: {len(doc.paragraphs)}")
        print()
        
        # Phân tích từng paragraph
        non_empty_paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                non_empty_paragraphs.append((i+1, text))
        
        print(f"📝 Paragraphs có nội dung: {len(non_empty_paragraphs)}")
        print()
        
        # Hiển thị nội dung
        print("📋 NỘI DUNG FILE:")
        print("-" * 80)
        
        for line_num, text in non_empty_paragraphs:
            print(f"Dòng {line_num}: {text}")
        
        print()
        
        # Phân tích loại nội dung
        headers = []
        questions = []
        other_content = []
        
        for line_num, text in non_empty_paragraphs:
            if re.match(r'^(Subject|Number|Lecturer|Date):', text, re.IGNORECASE):
                headers.append((line_num, text))
            elif re.match(r'^(QN|Câu|Question|Q\d+|^\d+)[:.]', text, re.IGNORECASE):
                questions.append((line_num, text))
            else:
                other_content.append((line_num, text))
        
        print("📊 PHÂN LOẠI NỘI DUNG:")
        print(f"📋 Headers: {len(headers)}")
        for line_num, text in headers:
            print(f"   Dòng {line_num}: {text}")
        
        print(f"\n❓ Câu hỏi: {len(questions)}")
        for line_num, text in questions:
            print(f"   Dòng {line_num}: {text}")
        
        print(f"\n📄 Nội dung khác: {len(other_content)}")
        for line_num, text in other_content:
            print(f"   Dòng {line_num}: {text}")
        
        # Đánh giá
        print("\n" + "=" * 80)
        print("📋 ĐÁNH GIÁ FILE")
        print("=" * 80)
        
        if len(questions) == 0:
            print("❌ VẤN ĐỀ: File không có câu hỏi!")
            print("💡 File này có vẻ chỉ chứa metadata hoặc header")
            print("🔧 Cần thêm nội dung câu hỏi thực tế")
            return False
        else:
            print(f"✅ File có {len(questions)} câu hỏi")
            return True
        
    except Exception as e:
        print(f"❌ Lỗi phân tích file: {e}")
        return False

def create_correct_template():
    """Tạo file template đúng định dạng"""
    print("\n" + "=" * 80)
    print("📝 TẠO FILE TEMPLATE ĐÚNG ĐỊNH DẠNG")
    print("=" * 80)
    
    try:
        doc = Document()
        
        # Header
        doc.add_paragraph("QUIZ TEMPLATE - ĐỊNH DẠNG ĐÚNG")
        doc.add_paragraph("Subject: ISC")
        doc.add_paragraph("Number of Quiz: 20")
        doc.add_paragraph("Lecturer: hungpd2")
        doc.add_paragraph("Date: dd-mm-yyyy")
        doc.add_paragraph("")
        
        # Câu hỏi 1 - có hình ảnh
        doc.add_paragraph("QN=1: See the figure and choose the right type of B2B E-Commerce [file:8435.jpg]")
        
        # Thêm hình ảnh mẫu
        try:
            from PIL import Image, ImageDraw
            
            # Tạo hình ảnh mẫu
            img = Image.new('RGB', (400, 300), color='white')
            draw = ImageDraw.Draw(img)
            
            # Vẽ diagram B2B
            draw.rectangle([50, 50, 350, 250], outline='blue', width=3)
            draw.text((200, 100), "An Exchange", fill='blue')
            draw.text((50, 30), "Sellers", fill='black')
            draw.text((350, 30), "Buyers", fill='black')
            draw.text((200, 30), "Services", fill='black')
            
            # Lưu tạm thời
            img_path = "temp_b2b_diagram.png"
            img.save(img_path)
            
            # Thêm vào document
            doc.add_picture(img_path, width=Inches(4))
            
            # Xóa file tạm
            os.remove(img_path)
            
        except ImportError:
            doc.add_paragraph("[B2B E-Commerce Diagram Placeholder]")
        
        # Đáp án
        doc.add_paragraph("a. Sell-side B2B")
        doc.add_paragraph("b. Electronic Exchange")
        doc.add_paragraph("c. Buy-side B2B")
        doc.add_paragraph("d. Supply Chain Improvements and Collaborative Commerce")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 2 - không có hình ảnh
        doc.add_paragraph("Câu 2: 2 + 2 = ?")
        doc.add_paragraph("A. 3")
        doc.add_paragraph("B. 4")
        doc.add_paragraph("C. 5")
        doc.add_paragraph("D. 6")
        doc.add_paragraph("Đáp án: B")
        doc.add_paragraph("Điểm: 0.5")
        doc.add_paragraph("Đơn vị: Toán học")
        doc.add_paragraph("")
        
        # Câu hỏi 3 - định dạng khác
        doc.add_paragraph("3. What is the capital of Vietnam?")
        doc.add_paragraph("a) Hanoi")
        doc.add_paragraph("b) Ho Chi Minh City")
        doc.add_paragraph("c) Da Nang")
        doc.add_paragraph("d) Hue")
        doc.add_paragraph("Correct: A")
        doc.add_paragraph("Mark: 1.0")
        doc.add_paragraph("Unit: Geography")
        
        # Lưu file
        output_file = "Quiz_Template_Correct_Format.docx"
        doc.save(output_file)
        
        print(f"✅ Đã tạo file template đúng: {output_file}")
        print("\n📋 NỘI DUNG FILE MẪU:")
        print("- Header với metadata")
        print("- 3 câu hỏi với các định dạng khác nhau")
        print("- 1 câu hỏi có hình ảnh")
        print("- 2 câu hỏi không có hình ảnh")
        print("- Đầy đủ đáp án và thông tin bổ sung")
        
        return output_file
        
    except Exception as e:
        print(f"❌ Lỗi tạo template: {e}")
        return None

def fix_existing_file(input_file, output_file):
    """Sửa file hiện tại bằng cách thêm nội dung mẫu"""
    print("\n" + "=" * 80)
    print("🔧 SỬA FILE HIỆN TẠI")
    print("=" * 80)
    
    try:
        # Đọc file hiện tại
        doc = Document(input_file)
        
        # Thêm nội dung câu hỏi mẫu
        doc.add_paragraph("")
        doc.add_paragraph("=== NỘI DUNG CÂU HỎI MẪU ===")
        doc.add_paragraph("")
        
        # Câu hỏi 1
        doc.add_paragraph("QN=1: See the figure and choose the right type of B2B E-Commerce [file:8435.jpg]")
        doc.add_paragraph("a. Sell-side B2B")
        doc.add_paragraph("b. Electronic Exchange")
        doc.add_paragraph("c. Buy-side B2B")
        doc.add_paragraph("d. Supply Chain Improvements and Collaborative Commerce")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 2
        doc.add_paragraph("Câu 2: 2 + 2 = ?")
        doc.add_paragraph("A. 3")
        doc.add_paragraph("B. 4")
        doc.add_paragraph("C. 5")
        doc.add_paragraph("D. 6")
        doc.add_paragraph("Đáp án: B")
        doc.add_paragraph("Điểm: 0.5")
        doc.add_paragraph("Đơn vị: Toán học")
        
        # Lưu file đã sửa
        doc.save(output_file)
        
        print(f"✅ Đã sửa file: {output_file}")
        print("📋 Đã thêm 2 câu hỏi mẫu vào file hiện tại")
        
        return output_file
        
    except Exception as e:
        print(f"❌ Lỗi sửa file: {e}")
        return None

def main():
    """Hàm main"""
    print("🚀 CÔNG CỤ KIỂM TRA VÀ SỬA FILE TEMPLATE")
    print("=" * 80)
    
    # Hỏi file cần kiểm tra
    file_path = input("Nhập đường dẫn file .docx cần kiểm tra: ").strip()
    
    if not file_path:
        print("❌ Vui lòng nhập đường dẫn file!")
        return
    
    # Phân tích file hiện tại
    has_questions = analyze_template_file(file_path)
    
    print("\n" + "=" * 80)
    print("🔧 GIẢI PHÁP")
    print("=" * 80)
    
    if has_questions:
        print("✅ File có câu hỏi hợp lệ!")
        print("💡 Bạn có thể sử dụng file này để test")
    else:
        print("❌ File không có câu hỏi hợp lệ!")
        print("\n🔧 CÁC GIẢI PHÁP:")
        print("1. Tạo file template mới đúng định dạng")
        print("2. Sửa file hiện tại bằng cách thêm nội dung mẫu")
        print("3. Sử dụng file mẫu có sẵn")
        
        choice = input("\nChọn giải pháp (1/2/3): ").strip()
        
        if choice == "1":
            # Tạo file mới
            new_file = create_correct_template()
            if new_file:
                print(f"\n✅ Đã tạo file mẫu: {new_file}")
                print("💡 Bạn có thể sử dụng file này để test chức năng import")
        
        elif choice == "2":
            # Sửa file hiện tại
            output_file = file_path.replace('.docx', '_fixed.docx')
            fixed_file = fix_existing_file(file_path, output_file)
            if fixed_file:
                print(f"\n✅ Đã sửa file: {fixed_file}")
                print("💡 File đã được thêm nội dung câu hỏi mẫu")
        
        elif choice == "3":
            # Tạo file mẫu
            sample_file = create_correct_template()
            if sample_file:
                print(f"\n✅ Đã tạo file mẫu: {sample_file}")
                print("💡 Sử dụng file này làm template cho các câu hỏi khác")
        
        else:
            print("❌ Lựa chọn không hợp lệ!")
    
    print("\n" + "=" * 80)
    print("🎉 HOÀN THÀNH")
    print("=" * 80)

if __name__ == "__main__":
    main() 