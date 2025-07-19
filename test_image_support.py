#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Test chức năng hỗ trợ hình ảnh trong file .docx
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from utils.docx_reader import DocxReader
from docx import Document
from docx.shared import Inches

def create_sample_file_with_images():
    """Tạo file mẫu có hình ảnh"""
    try:
        doc = Document()
        
        # Thêm tiêu đề
        doc.add_paragraph("SAMPLE QUESTION FILE WITH IMAGES")
        doc.add_paragraph("Format: QN=X: Question [file:image.jpg]")
        doc.add_paragraph("")
        
        # Câu hỏi 1 có hình ảnh
        doc.add_paragraph("QN=1: See the figure and choose the right type of B2B E-Commerce [file:8435.jpg]")
        
        # Thêm hình ảnh mẫu (nếu có)
        try:
            # Tạo hình ảnh đơn giản bằng PIL
            from PIL import Image, ImageDraw, ImageFont
            
            # Tạo hình ảnh mẫu
            img = Image.new('RGB', (400, 300), color='white')
            draw = ImageDraw.Draw(img)
            
            # Vẽ một hình đơn giản
            draw.rectangle([50, 50, 350, 250], outline='blue', width=3)
            draw.text((200, 150), "Sample Image", fill='black')
            
            # Lưu tạm thời
            img_path = "temp_sample_image.png"
            img.save(img_path)
            
            # Thêm vào document
            doc.add_picture(img_path, width=Inches(4))
            
            # Xóa file tạm
            os.remove(img_path)
            
        except ImportError:
            # Nếu không có PIL, tạo text thay thế
            doc.add_paragraph("[IMAGE PLACEHOLDER - Sample B2B E-Commerce Diagram]")
        
        # Đáp án
        doc.add_paragraph("a. Sell-side B2B")
        doc.add_paragraph("b. Electronic Exchange")
        doc.add_paragraph("c. Buy-side B2B")
        doc.add_paragraph("d. Supply Chain Improvements and Collaborative Commerce")
        doc.add_paragraph("ANSWER: B")
        doc.add_paragraph("MARK: 0.5")
        doc.add_paragraph("UNIT: Chapter1")
        doc.add_paragraph("MIX CHOICES: Yes")
        doc.add_paragraph("")
        
        # Câu hỏi 2 không có hình ảnh
        doc.add_paragraph("Câu 2: 2 + 2 = ?")
        doc.add_paragraph("A. 3")
        doc.add_paragraph("B. 4")
        doc.add_paragraph("C. 5")
        doc.add_paragraph("D. 6")
        doc.add_paragraph("Đáp án: B")
        doc.add_paragraph("Điểm: 0.5")
        doc.add_paragraph("Đơn vị: Toán học")
        
        # Lưu file
        sample_file = "sample_questions_with_images.docx"
        doc.save(sample_file)
        
        print(f"✅ Đã tạo file mẫu: {sample_file}")
        return sample_file
        
    except Exception as e:
        print(f"❌ Lỗi tạo file mẫu: {e}")
        return None

def test_image_extraction(file_path):
    """Test trích xuất hình ảnh"""
    print("=" * 80)
    print("🖼️ TEST TRÍCH XUẤT HÌNH ẢNH")
    print("=" * 80)
    
    if not os.path.exists(file_path):
        print(f"❌ File không tồn tại: {file_path}")
        return
    
    try:
        reader = DocxReader()
        
        # Test trích xuất hình ảnh
        print(f"📄 File: {os.path.basename(file_path)}")
        print()
        
        # Trích xuất hình ảnh
        images = reader.extract_images_from_docx(file_path)
        print(f"🖼️ Hình ảnh tìm thấy: {len(images)}")
        
        if images:
            print("\n📋 CHI TIẾT HÌNH ẢNH:")
            for i, img in enumerate(images, 1):
                print(f"  {i}. Tên: {img['name']}")
                print(f"     Kích thước: {len(img['data'])} bytes")
                print(f"     Paragraph: {img['paragraph_index']}")
                print()
        
        # Lưu hình ảnh
        saved_images = reader.save_images_to_folder(images)
        print(f"💾 Đã lưu {len(saved_images)} hình ảnh vào thư mục extracted_images/")
        
        if saved_images:
            print("\n📁 DANH SÁCH FILE ĐÃ LƯU:")
            for img in saved_images:
                print(f"  📄 {img['file_path']}")
        
        return len(images) > 0
        
    except Exception as e:
        print(f"❌ Lỗi test hình ảnh: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return False

def test_image_references(file_path):
    """Test xử lý tham chiếu hình ảnh"""
    print("\n" + "=" * 80)
    print("🔗 TEST XỬ LÝ THAM CHIẾU HÌNH ẢNH")
    print("=" * 80)
    
    try:
        reader = DocxReader()
        
        # Đọc file để tìm tham chiếu
        doc = Document(file_path)
        
        # Tìm tham chiếu hình ảnh trong text
        image_refs = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                refs = reader.extract_image_references_from_text(text)
                if refs:
                    image_refs.append({
                        'line': i + 1,
                        'text': text,
                        'refs': refs
                    })
        
        print(f"🔗 Tham chiếu hình ảnh tìm thấy: {len(image_refs)}")
        
        if image_refs:
            print("\n📝 CHI TIẾT THAM CHIẾU:")
            for item in image_refs:
                print(f"  Dòng {item['line']}: '{item['text']}'")
                for ref in item['refs']:
                    print(f"    → Tham chiếu: {ref}")
                print()
        
        return len(image_refs) > 0
        
    except Exception as e:
        print(f"❌ Lỗi test tham chiếu: {e}")
        return False

def test_full_processing(file_path):
    """Test xử lý hoàn chỉnh file có hình ảnh"""
    print("\n" + "=" * 80)
    print("🎯 TEST XỬ LÝ HOÀN CHỈNH")
    print("=" * 80)
    
    try:
        reader = DocxReader()
        
        # Test chi tiết
        success, result = reader.test_file_detailed(file_path)
        
        print("📊 KẾT QUẢ PHÂN TÍCH:")
        print("-" * 80)
        print(result)
        
        return success
        
    except Exception as e:
        print(f"❌ Lỗi test hoàn chỉnh: {e}")
        return False

def main():
    """Hàm main"""
    print("🚀 TEST HỖ TRỢ HÌNH ẢNH TRONG FILE .DOCX")
    print("=" * 80)
    
    # Tạo file mẫu có hình ảnh
    print("📝 Tạo file mẫu có hình ảnh...")
    sample_file = create_sample_file_with_images()
    
    if not sample_file:
        print("❌ Không thể tạo file mẫu!")
        return
    
    # Hỏi người dùng muốn test file nào
    print("\n📂 Chọn file để test:")
    print("1. File mẫu vừa tạo")
    print("2. File của bạn")
    
    choice = input("Nhập lựa chọn (1 hoặc 2): ").strip()
    
    if choice == "1":
        file_path = sample_file
    elif choice == "2":
        file_path = input("Nhập đường dẫn file .docx của bạn: ").strip()
        if not file_path:
            print("❌ Vui lòng nhập đường dẫn file!")
            return
    else:
        print("❌ Lựa chọn không hợp lệ!")
        return
    
    # Test các chức năng
    print(f"\n🔍 Bắt đầu test file: {file_path}")
    
    # Test trích xuất hình ảnh
    has_images = test_image_extraction(file_path)
    
    # Test tham chiếu hình ảnh
    has_refs = test_image_references(file_path)
    
    # Test xử lý hoàn chỉnh
    success = test_full_processing(file_path)
    
    # Tổng kết
    print("\n" + "=" * 80)
    print("📋 TỔNG KẾT KẾT QUẢ")
    print("=" * 80)
    print(f"🖼️ Có hình ảnh: {'✅ Có' if has_images else '❌ Không'}")
    print(f"🔗 Có tham chiếu: {'✅ Có' if has_refs else '❌ Không'}")
    print(f"✅ Xử lý thành công: {'✅ Có' if success else '❌ Không'}")
    
    if has_images or has_refs:
        print("\n💡 LƯU Ý:")
        print("- Hình ảnh được lưu trong thư mục extracted_images/")
        print("- Tham chiếu [file:filename] được xử lý tự động")
        print("- Bạn có thể sử dụng file này để test chức năng import")
    
    print("\n" + "=" * 80)
    print("🎉 HOÀN THÀNH TEST")
    print("=" * 80)

if __name__ == "__main__":
    main() 