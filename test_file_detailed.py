#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Test chi tiết file .docx để tìm hiểu tại sao không đọc được câu hỏi
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from utils.docx_reader import DocxReader
from docx import Document

def test_file_detailed(file_path):
    """Test file chi tiết"""
    print("=" * 80)
    print("🔍 TEST CHI TIẾT FILE .DOCX")
    print("=" * 80)
    
    if not os.path.exists(file_path):
        print(f"❌ File không tồn tại: {file_path}")
        return
    
    if not file_path.lower().endswith('.docx'):
        print(f"❌ File không phải định dạng .docx: {file_path}")
        return
    
    try:
        # Đọc file bằng python-docx
        doc = Document(file_path)
        print(f"📄 File: {os.path.basename(file_path)}")
        print(f"📊 Tổng paragraphs: {len(doc.paragraphs)}")
        print()
        
        # Phân tích từng paragraph
        non_empty_paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                non_empty_paragraphs.append((i+1, text))
        
        print(f"📝 Paragraphs có nội dung: {len(non_empty_paragraphs)}")
        print()
        
        # Hiển thị từng dòng
        print("📋 NỘI DUNG TỪNG DÒNG:")
        print("-" * 80)
        
        for line_num, text in non_empty_paragraphs:
            print(f"Dòng {line_num}: '{text}'")
        print()
        
        # Test với DocxReader
        reader = DocxReader()
        success, result = reader.test_file_detailed(file_path)
        
        print("🔍 KẾT QUẢ PHÂN TÍCH:")
        print("-" * 80)
        print(result)
        
        # Test từng hàm riêng lẻ
        print("\n🧪 TEST TỪNG HÀM NHẬN DIỆN:")
        print("-" * 80)
        
        for line_num, text in non_empty_paragraphs:
            print(f"\nDòng {line_num}: '{text}'")
            
            # Test câu hỏi
            is_question = reader._is_question_start(text)
            if is_question:
                question_num = reader._extract_question_number(text)
                print(f"  ✅ Là câu hỏi (số: {question_num})")
            else:
                print(f"  ❌ Không phải câu hỏi")
            
            # Test đáp án
            is_option = reader._is_option(text)
            if is_option:
                option_letter, option_text = reader._extract_option(text)
                print(f"  ✅ Là đáp án ({option_letter}: {option_text})")
            else:
                print(f"  ❌ Không phải đáp án")
            
            # Test đáp án đúng
            is_answer = reader._is_correct_answer(text)
            if is_answer:
                correct = reader._extract_correct_answer(text)
                print(f"  ✅ Là đáp án đúng ({correct})")
            else:
                print(f"  ❌ Không phải đáp án đúng")
            
            # Test điểm
            is_mark = reader._is_mark_info(text)
            if is_mark:
                mark = reader._extract_mark(text)
                print(f"  ✅ Là điểm ({mark})")
            else:
                print(f"  ❌ Không phải điểm")
            
            # Test đơn vị
            is_unit = reader._is_unit_info(text)
            if is_unit:
                unit = reader._extract_unit(text)
                print(f"  ✅ Là đơn vị ({unit})")
            else:
                print(f"  ❌ Không phải đơn vị")
            
            # Test trộn đáp án
            is_mix = reader._is_mix_choices_info(text)
            if is_mix:
                mix = reader._extract_mix_choices(text)
                print(f"  ✅ Là trộn đáp án ({mix})")
            else:
                print(f"  ❌ Không phải trộn đáp án")
        
        # Gợi ý sửa lỗi
        print("\n💡 GỢI Ý KHẮC PHỤC:")
        print("-" * 80)
        
        questions_found = sum(1 for _, text in non_empty_paragraphs if reader._is_question_start(text))
        options_found = sum(1 for _, text in non_empty_paragraphs if reader._is_option(text))
        answers_found = sum(1 for _, text in non_empty_paragraphs if reader._is_correct_answer(text))
        
        if questions_found == 0:
            print("❌ VẤN ĐỀ: Không tìm thấy câu hỏi nào!")
            print("\n🔧 CÁCH SỬA:")
            print("1. Đảm bảo câu hỏi bắt đầu bằng một trong các định dạng:")
            print("   - QN=1: Nội dung câu hỏi")
            print("   - Câu 1: Nội dung câu hỏi")
            print("   - 1. Nội dung câu hỏi")
            print("   - Q1: Nội dung câu hỏi")
            print("   - Question 1: Nội dung câu hỏi")
            print()
            print("2. Kiểm tra không có khoảng trắng thừa ở đầu dòng")
            print("3. Kiểm tra font chữ và encoding")
            print("4. Đảm bảo file được lưu đúng định dạng .docx")
        
        if options_found == 0:
            print("❌ VẤN ĐỀ: Không tìm thấy đáp án nào!")
            print("\n🔧 CÁCH SỬA:")
            print("1. Đảm bảo đáp án có định dạng:")
            print("   - a. Nội dung đáp án A")
            print("   - b. Nội dung đáp án B")
            print("   - c. Nội dung đáp án C")
            print("   - d. Nội dung đáp án D")
            print("   - A. Nội dung đáp án A")
            print("   - B. Nội dung đáp án B")
            print("   - C. Nội dung đáp án C")
            print("   - D. Nội dung đáp án D")
        
        if answers_found == 0:
            print("❌ VẤN ĐỀ: Không tìm thấy đáp án đúng nào!")
            print("\n🔧 CÁCH SỬA:")
            print("1. Đảm bảo đáp án đúng có định dạng:")
            print("   - ANSWER: A")
            print("   - Đáp án: B")
            print("   - Answer: C")
            print("   - Correct: D")
            print("   - Đúng: A")
        
        if questions_found > 0 and options_found > 0 and answers_found > 0:
            print("✅ File có vẻ hợp lệ!")
            print("Nếu vẫn không đọc được, có thể do:")
            print("1. Thứ tự các phần không đúng")
            print("2. Thiếu thông tin bắt buộc")
            print("3. Lỗi encoding")
        
    except Exception as e:
        print(f"❌ Lỗi đọc file: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")

def main():
    """Hàm main"""
    print("🚀 TEST CHI TIẾT FILE .DOCX")
    print("=" * 80)
    
    # Hỏi đường dẫn file
    file_path = input("Nhập đường dẫn file .docx: ").strip()
    
    if not file_path:
        print("❌ Vui lòng nhập đường dẫn file!")
        return
    
    # Test file
    test_file_detailed(file_path)
    
    print("\n" + "=" * 80)
    print("🎉 HOÀN THÀNH TEST")
    print("=" * 80)

if __name__ == "__main__":
    main() 