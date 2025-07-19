#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Test và khắc phục lỗi đọc file Word
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def test_imports():
    """Test các thư viện cần thiết"""
    print("=" * 60)
    print("TEST THƯ VIỆN CẦN THIẾT")
    print("=" * 60)
    
    try:
        # Test python-docx
        from docx import Document
        print("✅ python-docx imported successfully")
        
        # Test các module khác
        from utils.docx_reader import DocxReader
        print("✅ DocxReader imported successfully")
        
        from database.database_manager import DatabaseManager
        print("✅ DatabaseManager imported successfully")
        
        return True
        
    except ImportError as e:
        print(f"❌ Import error: {e}")
        print("🔧 Khắc phục:")
        print("   pip install python-docx")
        print("   pip install mysql-connector-python")
        return False
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        return False

def test_database_connection():
    """Test kết nối database"""
    print("\n" + "=" * 60)
    print("TEST KẾT NỐI DATABASE")
    print("=" * 60)
    
    try:
        from database.database_manager import DatabaseManager
        db = DatabaseManager()
        
        # Test query đơn giản
        result = db.execute_query("SELECT 1 as test")
        print("✅ Database connection successful")
        
        # Test bảng subjects
        subjects = db.execute_query("SELECT COUNT(*) as count FROM subjects")
        print(f"✅ Subjects table: {subjects[0]['count']} records")
        
        return True
        
    except Exception as e:
        print(f"❌ Database error: {e}")
        print("🔧 Khắc phục:")
        print("   1. Kiểm tra MySQL Server đã chạy chưa")
        print("   2. Kiểm tra config/database_config.py")
        print("   3. Đảm bảo database 'exam_bank' đã được tạo")
        return False

def test_docx_creation():
    """Test tạo file Word đơn giản"""
    print("\n" + "=" * 60)
    print("TEST TẠO FILE WORD")
    print("=" * 60)
    
    try:
        from docx import Document
        import tempfile
        
        # Tạo file Word đơn giản
        doc = Document()
        doc.add_paragraph("Test question")
        doc.add_paragraph("a. Option A")
        doc.add_paragraph("b. Option B")
        doc.add_paragraph("ANSWER: A")
        
        # Lưu file tạm
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            test_file_path = tmp_file.name
        
        print(f"✅ File Word created: {test_file_path}")
        
        # Test đọc file
        doc2 = Document(test_file_path)
        paragraphs = [p.text for p in doc2.paragraphs if p.text.strip()]
        print(f"✅ File read successfully: {len(paragraphs)} paragraphs")
        
        # Xóa file tạm
        os.unlink(test_file_path)
        print("✅ Test file deleted")
        
        return True
        
    except Exception as e:
        print(f"❌ File creation error: {e}")
        return False

def test_docx_reader():
    """Test DocxReader với file đơn giản"""
    print("\n" + "=" * 60)
    print("TEST DOCX READER")
    print("=" * 60)
    
    try:
        from utils.docx_reader import DocxReader
        from docx import Document
        import tempfile
        
        # Tạo file test
        doc = Document()
        doc.add_paragraph("QN=1: Test question?")
        doc.add_paragraph("a. Option A")
        doc.add_paragraph("b. Option B")
        doc.add_paragraph("c. Option C")
        doc.add_paragraph("d. Option D")
        doc.add_paragraph("ANSWER: A")
        
        # Lưu file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            test_file_path = tmp_file.name
        
        print(f"✅ Test file created: {test_file_path}")
        
        # Test DocxReader
        docx_reader = DocxReader()
        
        # Test với subject_id và creator_id giả
        success, message = docx_reader.read_docx_file(test_file_path, subject_id=1, creator_id=1)
        
        if success:
            print(f"✅ DocxReader successful: {message}")
        else:
            print(f"❌ DocxReader failed: {message}")
        
        # Xóa file
        os.unlink(test_file_path)
        
        return success
        
    except Exception as e:
        print(f"❌ DocxReader error: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_parsing_functions():
    """Test các hàm parse riêng lẻ"""
    print("\n" + "=" * 60)
    print("TEST PARSING FUNCTIONS")
    print("=" * 60)
    
    try:
        from utils.docx_reader import DocxReader
        docx_reader = DocxReader()
        
        # Test parse câu hỏi
        test_cases = [
            ("QN=1: Test", True),
            ("Câu 1: Test", True),
            ("1. Test", True),
            ("Test question", False),
        ]
        
        print("🔍 Testing question parsing:")
        for text, expected in test_cases:
            result = docx_reader._is_question_start(text)
            status = "✅" if result == expected else "❌"
            print(f"   {status} '{text}' -> {result} (expected: {expected})")
        
        # Test parse đáp án
        test_options = [
            ("a. Option A", True),
            ("A. Option A", True),
            ("a) Option A", True),
            ("Option A", False),
        ]
        
        print("\n🔍 Testing option parsing:")
        for text, expected in test_options:
            result = docx_reader._is_option(text)
            status = "✅" if result == expected else "❌"
            print(f"   {status} '{text}' -> {result} (expected: {expected})")
        
        # Test parse đáp án đúng
        test_answers = [
            ("ANSWER: A", True),
            ("Đáp án: B", True),
            ("Answer: C", True),
            ("Test answer", False),
        ]
        
        print("\n🔍 Testing answer parsing:")
        for text, expected in test_answers:
            result = docx_reader._is_correct_answer(text)
            status = "✅" if result == expected else "❌"
            print(f"   {status} '{text}' -> {result} (expected: {expected})")
        
        return True
        
    except Exception as e:
        print(f"❌ Parsing test error: {e}")
        return False

def create_sample_file():
    """Tạo file mẫu để test"""
    print("\n" + "=" * 60)
    print("TẠO FILE MẪU")
    print("=" * 60)
    
    try:
        from docx import Document
        
        # Tạo file mẫu
        doc = Document()
        
        # Thêm tiêu đề
        doc.add_paragraph("SAMPLE QUESTION FILE")
        doc.add_paragraph("Format: QN=X: Question")
        doc.add_paragraph("")
        
        # Câu hỏi 1
        doc.add_paragraph("QN=1: What is the capital of Vietnam?")
        doc.add_paragraph("a. Hanoi")
        doc.add_paragraph("b. Ho Chi Minh City")
        doc.add_paragraph("c. Da Nang")
        doc.add_paragraph("d. Hue")
        doc.add_paragraph("ANSWER: A")
        doc.add_paragraph("MARK: 1.0")
        doc.add_paragraph("UNIT: Geography")
        doc.add_paragraph("")
        
        # Câu hỏi 2
        doc.add_paragraph("Câu 2: 2 + 2 = ?")
        doc.add_paragraph("A. 3")
        doc.add_paragraph("B. 4")
        doc.add_paragraph("C. 5")
        doc.add_paragraph("D. 6")
        doc.add_paragraph("Đáp án: B")
        doc.add_paragraph("Điểm: 0.5")
        doc.add_paragraph("Đơn vị: Toán học")
        
        # Lưu file
        sample_file = "sample_questions.docx"
        doc.save(sample_file)
        print(f"✅ Sample file created: {sample_file}")
        print("📝 Bạn có thể sử dụng file này để test trong ứng dụng")
        
        return sample_file
        
    except Exception as e:
        print(f"❌ Error creating sample file: {e}")
        return None

def main():
    """Hàm main"""
    print("🔧 TEST VÀ KHẮC PHỤC LỖI ĐỌC FILE")
    print("=" * 60)
    
    all_tests_passed = True
    
    # Test imports
    if not test_imports():
        all_tests_passed = False
        print("\n⚠️ Vui lòng khắc phục lỗi import trước khi tiếp tục")
        return
    
    # Test database
    if not test_database_connection():
        all_tests_passed = False
        print("\n⚠️ Vui lòng khắc phục lỗi database trước khi tiếp tục")
        return
    
    # Test tạo file Word
    if not test_docx_creation():
        all_tests_passed = False
    
    # Test parsing functions
    if not test_parsing_functions():
        all_tests_passed = False
    
    # Test DocxReader
    if not test_docx_reader():
        all_tests_passed = False
    
    # Tạo file mẫu
    sample_file = create_sample_file()
    
    print("\n" + "=" * 60)
    if all_tests_passed:
        print("🎉 TẤT CẢ TEST THÀNH CÔNG!")
        print("✅ Hệ thống đọc file Word hoạt động bình thường")
        if sample_file:
            print(f"📁 File mẫu: {sample_file}")
            print("💡 Sử dụng file này để test trong ứng dụng chính")
    else:
        print("❌ CÓ LỖI XẢY RA!")
        print("🔧 Vui lòng khắc phục các lỗi trên trước khi sử dụng")
    
    print("\n📋 HƯỚNG DẪN KHẮC PHỤC:")
    print("1. Cài đặt thư viện: pip install python-docx mysql-connector-python")
    print("2. Kiểm tra MySQL Server đã chạy chưa")
    print("3. Kiểm tra config/database_config.py")
    print("4. Chạy lại test: python test_file_reading.py")

if __name__ == "__main__":
    main() 